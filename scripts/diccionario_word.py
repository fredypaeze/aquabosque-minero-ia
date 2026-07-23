"""Diccionario de Datos Maestro AquaBosque -> Word (.docx) landscape, editable."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERDE = "1B5E20"; VERDE2 = "2E7D32"; GRISF = "EEF4F1"

COLS = ["Página", "Componente Visual", "Nombre del Dato", "Definición Práctica",
        "Naturaleza", "Fuente Oficial", "Metodología / Construcción", "Interpretación (Rango)", "Medida"]
WIDTHS = [0.75, 1.0, 1.1, 1.95, 0.9, 0.85, 2.0, 1.05, 0.75]

ROWS = [
 # Inicio
 ["Inicio","Tarjeta KPI","Municipios analizados (1.122)","Cobertura del sistema: todo el país.","Metadato Territorial","DANE DIVIPOLA","Conteo de la base municipal oficial (1 fila/municipio).","Fijo: 1.122","municipios"],
 ["Inicio","Tarjeta KPI","Riesgo crítico (57)","Municipios de máxima prioridad de revisión.","Salida de Modelo (IA)","Cálculo propio","Conteo nivel Crítico (score ≥ percentil 95).","≈5% superior","municipios"],
 ["Inicio","Tarjeta KPI","Riesgo alto (112)","Municipios de atención prioritaria.","Salida de Modelo (IA)","Cálculo propio","Conteo score entre p85 y p95.","≈10%","municipios"],
 ["Inicio","Tarjeta KPI","Riesgo medio (280)","Municipios en seguimiento.","Salida de Modelo (IA)","Cálculo propio","Conteo score entre p60 y p85.","—","municipios"],
 # Mapa
 ["Mapa de riesgo","Mapa de coropletas","Nivel de riesgo por municipio","Colorea cada municipio según su prioridad.","Salida de Modelo (IA)","Cálculo propio + geom. DANE","Polígonos DIVIPOLA (GeoJSON) por riesgo_nivel.","Verde=Bajo…Rojo=Crítico","categórico (4)"],
 ["Mapa de riesgo","Mapa (Puntos)","Score de priorización (tamaño)","El punto crece con la prioridad.","Índice / Salida","Cálculo propio","Tamaño ∝ riesgo_score en el centroide.","0–1","adimensional"],
 ["Mapa de riesgo","Tarjeta KPI","Municipios en vista","Cuántos municipios se ven con los filtros.","Metadato Territorial","—","Conteo de filas filtradas.","conteo","municipios"],
 ["Mapa de riesgo","Tarjeta KPI","Crítico / Alto (en vista)","Críticos/altos en el filtro actual.","Salida de Modelo (IA)","Cálculo propio","Conteo por riesgo_nivel en el filtro.","conteo","municipios"],
 ["Mapa de riesgo","Tarjeta KPI","Score máx.","El puntaje más alto de la vista.","Índice / Salida","Cálculo propio","max(riesgo_score) de la vista.","0–1","adimensional"],
 ["Mapa de riesgo","Tabla superior","Top 5 (Mpio, Depto, Nivel, Score)","Los 5 más prioritarios de la vista.","Salida de Modelo (IA)","Cálculo propio","Orden descendente por riesgo_score.","0–1 / categórico","adim."],
 # Ranking
 ["Ranking","Gráfico de barras","Score de priorización","Ordena municipios de mayor a menor prioridad.","Índice / Salida","Cálculo propio","Barra = riesgo_score, color = riesgo_nivel.","0–1","adimensional"],
 ["Ranking","Deslizador","Top N","Cuántos municipios mostrar.","Metadato","—","Control de UI (10–200).","10–200","conteo"],
 ["Ranking","Tabla","Minero / Deforest. / Hídrico / Sensib.","Los 4 índices que componen el puntaje.","Índice Normalizado","ANM / IDEAM / IDEAM / RUNAP+PDET","Ver Anexo A (log1p + MinMax).","0–1","adimensional"],
 ["Ranking","Botón","Descargar ranking (CSV)","Exporta la tabla completa.","—","—","Descarga del dataframe a CSV.","—","archivo"],
 # Ficha
 ["Ficha territorial","Tarjeta KPI","Nivel de riesgo","La prioridad del municipio en palabras.","Salida de Modelo (IA)","Cálculo propio","Clasificación por cuantil del score.","Bajo/Medio/Alto/Crítico","categórico"],
 ["Ficha territorial","Tarjeta KPI","Score de priorización","El puntaje técnico exacto.","Índice / Salida","Cálculo propio","Suma ponderada de 5 índices.","0–1","adimensional"],
 ["Ficha territorial","Tarjeta KPI","Predicción del modelo","Nivel que asigna la IA (XGBoost).","Salida de Modelo (IA)","Cálculo propio","argmax de XGBoost multiclase (15 vars).","Bajo…Crítico","categórico"],
 ["Ficha territorial","Tarjeta KPI","Confianza","Qué tan seguro está el modelo.","Salida de Modelo (IA)","Cálculo propio","max(predict_proba) de XGBoost.","0–100%","%"],
 ["Ficha territorial","Gráfico de radar","Perfil por dimensión","Qué factor domina en el municipio.","Índice Normalizado","ANM/IDEAM/IDEAM/RUNAP","4 índices 0–1 en ejes de radar.","0–1 por eje","adimensional"],
 ["Ficha territorial","Tarjeta 'Factor dominante'","Mayor factor de priorización","El índice más alto del municipio.","Índice Normalizado","(según dimensión)","argmax de los índices del municipio.","0–1","adimensional"],
 # Explicabilidad
 ["Explicabilidad","Gráfico de barras SHAP","Importancia SHAP (media |valor|)","Qué variables pesan más en el modelo, en promedio.","Salida de Modelo (IA)","Cálculo propio (TreeExplainer)","Media |SHAP| por variable sobre 4 clases y todos los municipios.","≥ 0","adim. (media |SHAP|)"],
 ["Explicabilidad","Tarjeta KPI","Accuracy (89.0%)","Aciertos del modelo en la prueba.","Salida de Modelo (IA)","Cálculo propio","25% test. Alta por construcción (etiqueta=fórmula); no es el mérito.","0–100%","%"],
 ["Explicabilidad","Tarjeta KPI","Línea base clase mayoritaria (60.1%)","Qué acertaría 'adivinar lo más común'.","Salida de Modelo (IA)","Cálculo propio","Frecuencia de la clase mayoritaria en test.","0–100%","%"],
 ["Explicabilidad","Tarjeta KPI","F1-macro (0.78)","Calidad equilibrada entre las 4 clases.","Salida de Modelo (IA)","Cálculo propio","Promedio no ponderado del F1 de las 4 clases.","0–1","adimensional"],
 # Datos abiertos
 ["Datos abiertos","Tabla de fuentes","Dimensión/Entidad/URL/Registros/Clave/Limitación","Ficha de cada fuente: qué mide, quién y cómo se cruza.","Metadato","ANM,IDEAM,DANE,RUNAP,PDET","Documenta origen y cruce (código DANE vs centroide).","—","según fila"],
 # Monitoreo satelital
 ["Monitoreo satelital","Tarjeta KPI","🔥 Focos (7 días)","Puntos de calor detectados por satélite esta semana.","Fuente Cruda","NASA FIRMS (VIIRS+MODIS)","Suma de detecciones asignadas a municipios (7 días).","conteo","focos"],
 ["Monitoreo satelital","Tarjeta KPI","Municipios con fuego","Municipios con ≥1 foco activo.","Fuente Cruda","NASA FIRMS","Conteo con focos_7d > 0.","conteo","municipios"],
 ["Monitoreo satelital","Tarjeta KPI","⚠️ Prioridad máxima","Modelo prioriza Y satélite confirma fuego hoy.","Salida IA + Cruda","Cálculo propio + FIRMS","Nivel∈{Alto,Crítico} y focos_7d>0.","conteo","municipios"],
 ["Monitoreo satelital","Tarjeta KPI","🆕 Actividad nueva","Fuego intenso donde el índice histórico no marcaba.","Salida IA + Cruda","Cálculo propio + FIRMS","frp_total>200 y nivel∈{Bajo,Medio}.","conteo","municipios"],
 ["Monitoreo satelital","Mapa de coropletas","Focos por municipio","Intensidad de fuego reciente por municipio.","Fuente Cruda","NASA FIRMS","Coropleta por focos_7d.","0 → alto","focos"],
 ["Monitoreo satelital","Tabla","Focos 7d","Número de focos del municipio (7 días).","Fuente Cruda","NASA FIRMS","Conteo de detecciones por polígono municipal.","≥ 0","focos"],
 ["Monitoreo satelital","Tabla","FRP (frp_total)","Intensidad acumulada del fuego (qué tan fuerte quema).","Fuente Cruda","NASA FIRMS","Suma de la Potencia Radiativa del Fuego (Fire Radiative Power).","≥ 0","MW (megavatios)"],
 # Asistente
 ["Asistente IA","Respuesta del asistente","Respuesta aterrizada + fuentes","Explica en lenguaje natural, solo con la evidencia.","Salida de Modelo (IA)","LLM local (Ministerio)","Recuperación con bge-m3 + generación con LLM local; cita fuente, no inventa.","texto","—"],
]

ANEXO_A_HDR = ["Índice","Definición Práctica","Fuente Oficial","Metodología / Fórmula","Rango","Medida"]
ANEXO_A = [
 ["idx_minero","Presión minera formal del municipio.","ANM (RUCOM + volumen)","minero_raw = titulos + volumen/1000 → log1p + MinMax. Cruce por CÓDIGO DANE (exacto).","0–1 (1=mayor)","adimensional"],
 ["idx_deforestacion","Pérdida de bosque registrada.","IDEAM / SMByC","deforestacion_ha → log1p + MinMax. Cruce POR NOMBRE; sin registro=0.","0–1 (1=más)","adimensional"],
 ["idx_fuego","Señal satelital de quema reciente (dinámico).","NASA FIRMS","log1p(frp_total)/log1p(máx), acotado [0,1]. Refresco diario.","0–1 (1=intenso)","adimensional"],
 ["idx_hidrico","Afectación de la calidad del agua.","IDEAM (DHIME · ICA)","1 − ICA (estación <50 km del centroide, haversine); sin estación=0 y marca hidrico_sin_dato.","0–1 (1=degradada)","adimensional"],
 ["idx_sensibilidad","Valor ambiental/social a proteger.","RUNAP + PDET","runap_hectareas → log1p+MinMax; +0.25 si PDET. RUNAP por centroide; PDET por DANE.","0–1 (1=mayor)","adimensional"],
]

ANEXO_B_HDR = ["Variable cruda","Definición","Fuente","Cruce","Medida"]
ANEXO_B = [
 ["mineria_titulos","Nº de títulos/explotadores mineros formales","ANM RUCOM (42ha-fhvj)","código DANE","títulos (conteo)"],
 ["mineria_volumen","Volumen de explotación (último año)","ANM (r85m-vv6c)","código DANE","heterogénea (t/m³ según mineral)"],
 ["mineria_regalias","Regalías pagadas","ANM","código DANE","COP (pesos)"],
 ["deforestacion_ha","Hectáreas deforestadas (último año)","IDEAM / SMByC","nombre de municipio","ha (hectáreas)"],
 ["agua_ica_medio","ICA medio de estaciones cercanas","IDEAM DHIME (ica5)","centroide, <50 km","adimensional (0–1; 1=buena)"],
 ["agua_estaciones","Nº de estaciones ICA asignadas","IDEAM DHIME","centroide","estaciones (conteo)"],
 ["runap_areas","Nº de áreas protegidas cercanas","RUNAP","centroide","áreas (conteo)"],
 ["runap_hectareas","Hectáreas protegidas cercanas","RUNAP","centroide","ha (hectáreas)"],
 ["es_pdet","¿Municipio PDET (posconflicto)?","DANE / PDET","código DANE","binario (0/1)"],
 ["focos_7d","Focos de calor activos (7 días)","NASA FIRMS","point-in-polygon","focos (conteo)"],
 ["frp_total","Potencia radiativa acumulada del fuego","NASA FIRMS","point-in-polygon","MW (megavatios)"],
]

INCERT = [
 "mineria_volumen — unidad heterogénea (t/m³ según mineral): por eso se usa log+MinMax (relativo), no valor absoluto.",
 "idx_hidrico — solo ~71 municipios con estación ICA a <50 km; 0 significa 'sin dato observado', NO 'agua sana'.",
 "deforestacion_ha — cruce por nombre de municipio (posibles omisiones por homónimos/tildes). Sin registro = 0.",
 "idx_fuego / Sentinel-2 (U-Net): la detección profunda por imagen es prueba de capacidad; la señal operativa en la app es FIRMS (focos térmicos).",
 "Accuracy alta ≠ mérito predictivo: la etiqueta es una fórmula; lo defendible es la priorización interpretable (SHAP).",
]


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, size=8.5, bold=False, color="222A26", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    r.font.name = "Calibri"; r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None, hdr_fill=VERDE):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; set_cell(c, h, 8.5, True, "FFFFFF"); shade(c, hdr_fill)
        if widths and j < len(widths): c.width = Inches(widths[j])
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], str(val), 8.3, bold=(j == 0))
            if i % 2 == 1: shade(cells[j], GRISF)
            if widths and j < len(widths): cells[j].width = Inches(widths[j])
    return t


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.5))

    h = doc.add_heading("", level=0)
    r = h.add_run("Diccionario de Datos Maestro — AquaBosque Minero IA")
    r.font.size = Pt(20); r.font.color.rgb = RGBColor.from_string(VERDE); r.font.name = "Calibri"
    p = doc.add_paragraph()
    rp = p.add_run("Cheat sheet de defensa ante jurado · basado en el código real del repositorio (build_master, build_target, train, firms_signal) y los datos vigentes.")
    rp.font.size = Pt(10); rp.font.italic = True; rp.font.color.rgb = RGBColor.from_string("5B6B62")
    nota = doc.add_paragraph()
    rn = nota.add_run("Nota de honestidad: la etiqueta de riesgo es una fórmula técnica de priorización, no una medición de daño. Ningún dato prueba causalidad ni ilegalidad.")
    rn.font.size = Pt(9.5); rn.font.bold = True; rn.font.color.rgb = RGBColor.from_string("7C4A03")

    doc.add_heading("Tabla maestra — por página de la app", level=1)
    add_table(doc, COLS, ROWS, WIDTHS)

    doc.add_heading("Anexo A · Índices normalizados (el corazón del score)", level=1)
    add_table(doc, ANEXO_A_HDR, ANEXO_A, [1.0, 2.2, 1.3, 3.4, 1.0, 1.1], hdr_fill=VERDE2)
    f = doc.add_paragraph()
    rf = f.add_run("Fórmula:  score = 0.30·minero + 0.25·deforestación + 0.15·fuego + 0.20·hídrico + 0.10·sensibilidad  (≈0–0.5).   "
                   "Nivel por cuantiles: Crítico ≥ p95 · Alto ≥ p85 · Medio ≥ p60 · Bajo < p60.   "
                   "Modelo: XGBoost multiclase · 15 variables · 300 árboles · prof. 4 · lr 0.08.")
    rf.font.size = Pt(9.5); rf.font.bold = True; rf.font.color.rgb = RGBColor.from_string("123A5C")

    doc.add_heading("Anexo B · Fuentes crudas y unidades", level=1)
    add_table(doc, ANEXO_B_HDR, ANEXO_B, [1.4, 2.6, 1.6, 1.5, 2.0], hdr_fill=VERDE2)

    doc.add_heading("Anexo C · Incertidumbres declaradas (cero alucinaciones)", level=1)
    for it in INCERT:
        pp = doc.add_paragraph(style="List Bullet")
        rr = pp.add_run(it); rr.font.size = Pt(10); rr.font.color.rgb = RGBColor.from_string("222A26")

    out = Path(__file__).resolve().parents[1] / "outputs" / "jurado_2026" / "DICCIONARIO_DATOS_MAESTRO.docx"
    doc.save(out); print("OK:", out)


if __name__ == "__main__":
    build()
